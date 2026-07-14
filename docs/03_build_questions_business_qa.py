"""
RankUp Education — Questions Business QA regenerator.

Single source of truth for:
  - docs/03_RankUp_Questions_Business_QA.html
  - docs/03_RankUp_Questions_Business_QA.docx

Run (from repo root, Python 3.12+ with python-docx):
  python docs/03_build_questions_business_qa.py
"""

from __future__ import annotations

import html
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parent
HTML_OUTPUT = DOCS_DIR / "03_RankUp_Questions_Business_QA.html"
DOCX_OUTPUT = DOCS_DIR / "03_RankUp_Questions_Business_QA.docx"

# ---------------------------------------------------------------------------
# Content — TARGET business rules (await implementation command)
# ---------------------------------------------------------------------------

META = {
    "title": "RankUp Education — Questions Business & QA Guide",
    "subtitle": "Question-bank permissions, type scope (incl. Fill in the Blanks), data model, Web Excel import, and gap vs current code.",
    "prepared_for": "RankUp Education workspace",
    "prepared_on": "15 Jul 2026",
    "status": "TARGET RULES — Phase 0–3 done (15 Jul 2026)",
    "roles": "PortalAdmin · SchoolAdmin · CampusAdmin · Teacher · Parent (Student: quizzes only)",
}

# --- Lookups (from user / DB) ---

DIFFICULTY_ROWS = [
    ("2001", "Easy", "DifficultyLevel", "1"),
    ("2002", "Medium", "DifficultyLevel", "2"),
    ("2003", "Hard", "DifficultyLevel", "3"),
]

STATUS_ROWS = [
    ("110", "Draft", "QuestionStatus", "Editable by owner; not yet submitted for review (optional use)."),
    ("111", "PendingReview", "QuestionStatus", "Submitted; waiting for PortalAdmin approval."),
    ("112", "Approved", "QuestionStatus", "Visible to everyone; locked for non–PortalAdmin edit/delete."),
    ("113", "Rejected", "QuestionStatus", "PortalAdmin reject (reason required). Owner may edit+resubmit to PendingReview, or delete."),
    ("114", "Archived", "QuestionStatus", "Retired / hidden from normal bank lists (PortalAdmin)."),
]

TYPE_ROWS = [
    ("100", "Single Choice", "QuestionType", "1", "NOW"),
    ("101", "Multiple Choice", "QuestionType", "2", "NOW"),
    ("102", "True/False", "QuestionType", "3", "NOW"),
    ("103", "Fill in the Blanks", "QuestionType", "4", "NOW"),
    ("104", "Descriptive", "QuestionType", "5", "NOT NOW"),
]

# Future types (no lookup IDs until implementation) — documented for roadmap only.
FUTURE_TYPE_ROWS = [
    ("—", "Matching", "Future", "Not now — implement later"),
    ("—", "Ordering", "Future", "Not now — implement later"),
    ("—", "Image-based", "Future", "Not now — implement later"),
    ("—", "Audio-based", "Future", "Not now — implement later"),
    ("—", "Mathematical working", "Future", "Not now — implement later"),
    ("—", "Coding answer", "Future", "Not now — implement later"),
    ("—", "Case-study", "Future", "Not now — implement later"),
]

TYPE_RULE_ROWS = [
    (
        "Single Choice (100)",
        "NOW",
        "question_options",
        "≥2 options; exactly one option must have IsCorrect = true.",
    ),
    (
        "Multiple Choice (101)",
        "NOW",
        "question_options",
        "≥2 options; at least one option must have IsCorrect = true; multiple may be correct.",
    ),
    (
        "True/False (102)",
        "NOW",
        "question_options",
        "Exactly two options (True / False); exactly one IsCorrect = true.",
    ),
    (
        "Fill in the Blanks (103)",
        "NOW",
        "question_accepted_answers",
        "≥1 accepted answer. Optional AI initial review + teacher review on attempts. "
        "Model / accepted answers must NOT be shown to the student before submission.",
    ),
    (
        "Descriptive (104)",
        "NOT NOW",
        "question_accepted_answers",
        "Same review/visibility rules as Fill when built. Hide create/import until ready.",
    ),
]

QUESTION_FIELD_ROWS = [
    ("Id", "PK"),
    ("QuestionText", "Stem / prompt"),
    ("QuestionType", "Lookup: Single Choice, Multiple Choice, True/False, Fill in the Blanks, Descriptive"),
    ("Class", "Lookup / ClassId"),
    ("Subject", "Lookup / SubjectId"),
    ("Topic", "Lookup / TopicId (optional)"),
    ("DifficultyLevel", "Easy / Medium / Hard"),
    ("Explanation", "Shown after attempt per quiz rules; bank detail for manage roles"),
    ("Hint", "Optional hint"),
    ("EstimatedTimeSeconds", "Estimated time"),
    ("Marks", "Marks / points"),
    ("IsActive", "Activate / Deactivate (PortalAdmin only)"),
    ("Status", "Draft, PendingReview, Approved, Rejected, Archived"),
    ("CreatedBy", "Owner"),
    ("IsAiApproved", "Optional system/AI quality flag — not a second bank gate for quiz eligibility"),
    ("ApprovedBy", "PortalAdmin who approved (human bank approval)"),
    ("CreatedDate", "Created timestamp"),
    ("ModifiedDate", "Modified timestamp"),
]

OPTION_FIELD_ROWS = [
    ("Id", "PK"),
    ("QuestionId", "FK → questions"),
    ("OptionText", "Choice text"),
    ("OptionImageUrl", "Optional image (future image-based use; column kept)"),
    ("IsCorrect", "Whether this option is a correct choice"),
    ("Explanation", "Optional per-option explanation"),
    ("IsActive", "Soft-disable option without delete"),
]

ACCEPTED_ANSWER_FIELD_ROWS = [
    ("Id", "PK"),
    ("QuestionId", "FK → questions"),
    ("AnswerText", "Canonical accepted answer text"),
    ("IsCaseSensitive", "Match case when scoring"),
    ("AllowPartialMatch", "Allow substring / partial credit rules"),
    ("NormalizedAnswer", "Normalized form for matching (e.g. lowercase trim)"),
    ("MinimumLength", "Min student answer length"),
    ("MaximumLength", "Max student answer length"),
    ("AllowAIReview", "Attempt may get AI initial review (Fill / Descriptive)"),
    ("AllowTeacherReview", "Attempt may require teacher review (Fill / Descriptive)"),
]

# --- Target rules matrix ---

PERMISSION_ROWS = [
    (
        "Create question",
        "Yes (any)",
        "Yes",
        "Yes",
        "Yes",
        "Yes",
        "No",
    ),
    (
        "Update own question (not Approved)",
        "Yes (any)",
        "Own only",
        "Own only",
        "Own only",
        "Own only",
        "No",
    ),
    (
        "Update Rejected → resubmit PendingReview",
        "Yes (any)",
        "Own only",
        "Own only",
        "Own only",
        "Own only",
        "No",
    ),
    (
        "Update Approved question",
        "Yes (any)",
        "No",
        "No",
        "No",
        "No",
        "No",
    ),
    (
        "Delete own (Draft / PendingReview / Rejected)",
        "Yes (any)",
        "Own only",
        "Own only",
        "Own only",
        "Own only",
        "No",
    ),
    (
        "Delete Approved question",
        "Yes (any)",
        "No",
        "No",
        "No",
        "No",
        "No",
    ),
    (
        "Activate / Deactivate / Archive",
        "Yes — ONLY PortalAdmin",
        "No",
        "No",
        "No",
        "No",
        "No",
    ),
    (
        "Approve",
        "Yes — ONLY PortalAdmin",
        "No",
        "No",
        "No",
        "No",
        "No",
    ),
    (
        "Reject (reason required)",
        "Yes — ONLY PortalAdmin",
        "No",
        "No",
        "No",
        "No",
        "No",
    ),
    (
        "View all Approved (+ answers / details)",
        "Yes",
        "Yes",
        "Yes",
        "Yes",
        "Yes",
        "No (quizzes only)",
    ),
    (
        "View own Draft / PendingReview / Rejected",
        "Yes (all)",
        "Own",
        "Own",
        "Own",
        "Own",
        "No",
    ),
    (
        "Import from Excel (Web only)",
        "Yes",
        "Yes",
        "Yes",
        "Yes",
        "Yes",
        "No",
    ),
    (
        "Import from Excel (Mobile)",
        "No",
        "No",
        "No",
        "No",
        "No",
        "No",
    ),
]

PERMISSION_HEADERS = [
    "Action",
    "PortalAdmin",
    "SchoolAdmin",
    "CampusAdmin",
    "Teacher",
    "Parent",
    "Student",
]

TARGET_RULES = [
    (
        "1",
        "PortalAdmin full control",
        "PortalAdmin (Admin Portal) can Create, Update, Delete any question; Activate / Deactivate / Archive; Approve; and Reject (rejection reason required).",
    ),
    (
        "2",
        "Create",
        "SchoolAdmin, CampusAdmin, Teacher, and Parent can create questions.",
    ),
    (
        "3",
        "Update own (pre-approval)",
        "SchoolAdmin, CampusAdmin, Teacher, and Parent can update their own questions only while status is not Approved (Draft, PendingReview, Rejected).",
    ),
    (
        "4",
        "Delete own (pre-approval)",
        "SchoolAdmin, CampusAdmin, Teacher, and Parent can delete their own questions only while status is Draft, PendingReview, or Rejected (not Approved).",
    ),
    (
        "5",
        "Delete after Approved",
        "Once Approved, only PortalAdmin may delete. Non–PortalAdmin cannot delete Approved questions (even their own).",
    ),
    (
        "6",
        "Activate / Deactivate / Archive",
        "Only PortalAdmin can Activate, Deactivate, or Archive questions. No other role.",
    ),
    (
        "7",
        "View Approved (creators + admins)",
        "PortalAdmin, SchoolAdmin, CampusAdmin, Teacher, and Parent can view all Approved questions with options/answers and full details. Students cannot open the question bank at all — they only see questions inside quizzes.",
    ),
    (
        "8",
        "Approve / Reject = PortalAdmin only",
        "Only PortalAdmin approves or rejects. Reject requires a non-empty rejection reason. After Reject: owner may edit and resubmit to PendingReview, or delete the question.",
    ),
    (
        "9",
        "Excel import (Web only)",
        "Import from Excel is available on Web for create-capable roles (PortalAdmin, SchoolAdmin, CampusAdmin, Teacher, Parent). Not available on the Mobile app.",
    ),
]

LIFECYCLE = [
    ("Create / Import (Web)", "Draft or PendingReview", "Owner chooses save-as-draft vs submit; import must never set Approved."),
    ("Submit for review", "PendingReview", "Owner submits; appears on PortalAdmin approval queue."),
    ("PortalAdmin Approve", "Approved", "Visible in bank to non-Student manage roles; locked for non–PortalAdmin edit/delete."),
    (
        "PortalAdmin Reject",
        "Rejected",
        "Rejection reason REQUIRED. Owner may (a) edit + resubmit → PendingReview, or (b) delete the question.",
    ),
    ("PortalAdmin Archive", "Archived", "Only PortalAdmin. Hidden from normal Approved lists."),
    ("PortalAdmin Activate / Deactivate", "IsActive flag", "Only PortalAdmin. Separate from status where product keeps both."),
]

IMPORT_COLUMNS = [
    ("QuestionText", "Required", "Stem / prompt text"),
    (
        "QuestionType",
        "Required",
        "NOW: Single Choice (100), Multiple Choice (101), True/False (102), Fill in the Blanks (103). "
        "NOT NOW: Descriptive (104) — reject or skip until enabled. Future types: not allowed.",
    ),
    ("Class", "Required", "Lookup name or ID"),
    ("Subject", "Required", "Lookup name or ID"),
    ("Topic", "Optional", "Lookup name or ID"),
    ("DifficultyLevel", "Required", "Easy (2001) / Medium (2002) / Hard (2003) — name or ID"),
    ("Marks", "Required", "Positive number"),
    ("EstimatedTimeSeconds", "Optional", "Default if blank"),
    ("Hint", "Optional", ""),
    ("Explanation", "Optional", ""),
    (
        "Option1…OptionN",
        "Choice types",
        "Required for Single Choice / Multiple Choice / True/False. Mark correct with IsCorrect1… or CorrectOption index.",
    ),
    (
        "AcceptedAnswer1…N",
        "Fill in the Blanks",
        "One or more accepted answers for type 103. Optional columns: IsCaseSensitive, AllowPartialMatch, MinLength, MaxLength, AllowAIReview, AllowTeacherReview.",
    ),
    ("Status", "Optional", "Draft or PendingReview only on import; never Approved via import"),
]

GAP_ROWS = [
    (
        "Web list Edit/Delete buttons",
        "QuestionsPage gated: View always; Edit/Delete for PortalAdmin or owner + Draft/PendingReview/Rejected",
        "Gate like Detail page (PortalAdmin or owner + Draft/PendingReview/Rejected)",
        "DONE",
    ),
    (
        "Web Save as Draft",
        "Create form: Save as Draft (SubmitForReview=false) vs Submit for review & Add",
        "Create form: Save as Draft vs Submit for review",
        "DONE",
    ),
    (
        "Canonical status IDs 110–114",
        "Writes prefer IDs 110–114; seed/remap on API start; canonical names Draft/PendingReview/Approved/Rejected/Archived (legacy aliases still readable)",
        "Prefer fixed IDs / names Draft, PendingReview, Approved, Rejected, Archived only",
        "DONE",
    ),
    (
        "IsAiApproved leftover",
        "Dual gate removed; Mobile/Web no AI copy; API docs note IsAiApproved is legacy (column kept)",
        "Clean messaging; optional deprecate field",
        "DONE",
    ),
    (
        "Fill in the Blanks storage",
        "Fill create/update/import use question_accepted_answers; legacy options migrated on API start; attempt omits answers pre-submit",
        "Use question_accepted_answers only for type 103; enforce ≥1 answer; wire create/update/import",
        "DONE",
    ),
    (
        "Type validation rules",
        "API + Web enforce Single=exactly 1 correct; Multi=≥1; T/F=exactly 2 + 1 correct; Fill=≥1 accepted; Descriptive rejected",
        "Enforce Single=exactly 1 correct; Multi=≥1; T/F=exactly 2 opts + 1 correct; Fill=≥1 accepted answer",
        "DONE",
    ),
    (
        "Hide Descriptive + future types in UI",
        "Create/import UI and ResolveQuestionType offer NOW types only (100–103); Descriptive rejected",
        "Create/import: NOW types only (100–103); hide Descriptive until enabled",
        "DONE",
    ),
    (
        "Attempt answer visibility (Fill)",
        "Student start-attempt returns empty options for Fill; scoring uses accepted answers server-side",
        "Student attempt APIs omit accepted/model answers until after submit (or review)",
        "DONE",
    ),
    (
        "Mobile question bank",
        "Router + page gate: manage roles only; Student redirected; no AI approved copy",
        "Manage roles only; Student must not open bank route",
        "DONE",
    ),
    (
        "Import dry-run UI",
        "Web: Dry run Excel + full error list + Confirm import after clean dry run",
        "Optional Dry run button + full error list",
        "DONE",
    ),
    (
        "Delete while quiz-linked",
        "Hard block for everyone (correct)",
        "Optional guided unlink then delete for PortalAdmin",
        "OPTIONAL",
    ),
    (
        "Approve/Reject audit trail",
        "Not logged beyond question fields",
        "Actor + timestamp log for Approve/Reject/Archive",
        "OPTIONAL",
    ),
    (
        "Core permission matrix",
        "Pre-fix: SchoolAdmin approve, no CampusAdmin, AI gate, no import",
        "PortalAdmin-only approve/lifecycle; CampusAdmin create; owner pre-approval edit/delete; Web Excel import; single Approve = quiz-ready",
        "DONE",
    ),
]

SCENARIOS = [
    (
        "Q-01",
        "PortalAdmin CRUD + Approve + Activate",
        [
            "PortalAdmin creates a question → Approve → status Approved.",
            "PortalAdmin Activate / Deactivate / Archive → allowed; other roles → forbidden.",
            "PortalAdmin updates/deletes any question → allowed (respect quiz-link policy if still enforced).",
        ],
        "Full control for PortalAdmin including Activate/Deactivate/Archive.",
    ),
    (
        "Q-02",
        "Teacher creates and edits own Pending",
        [
            "Teacher creates question → PendingReview (or Draft then submit).",
            "Teacher updates own Pending question → allowed.",
            "Teacher deletes own Pending question → allowed.",
        ],
        "Owner can manage pre-Approved own items.",
    ),
    (
        "Q-03",
        "Teacher blocked on Approved",
        [
            "PortalAdmin approves Teacher’s question.",
            "Teacher tries Update → forbidden.",
            "Teacher tries Delete → forbidden.",
            "Teacher tries Activate/Deactivate → forbidden.",
        ],
        "Approved is locked for non–PortalAdmin; lifecycle flags are PortalAdmin-only.",
    ),
    (
        "Q-04",
        "CampusAdmin create",
        [
            "Login as CampusAdmin → Create question succeeds.",
            "CampusAdmin sees own Pending; can edit/delete until Approved.",
        ],
        "CampusAdmin is a first-class creator.",
    ),
    (
        "Q-05",
        "SchoolAdmin cannot approve or reject",
        [
            "SchoolAdmin opens pending queue → no Approve/Reject (or API 403).",
            "PortalAdmin Approves → status Approved.",
        ],
        "Approval/rejection is PortalAdmin-only.",
    ),
    (
        "Q-06",
        "Approved visible to creators; Student blocked from bank",
        [
            "Create + Approve a question with options.",
            "As Teacher / Parent / SchoolAdmin / CampusAdmin: open Approved → see text, options, correct answers, explanation.",
            "As Student: no Questions bank menu; API list/get bank → forbidden. Student still answers questions inside a quiz attempt.",
        ],
        "Bank read for manage roles; Students only via quizzes.",
    ),
    (
        "Q-07",
        "Excel import — Web only",
        [
            "On Web, creator clicks Import → uploads valid .xlsx.",
            "Rows create questions owned by importer; status Draft or PendingReview (never Approved).",
            "Invalid rows reported with row errors.",
            "On Mobile: no Import button / no import API usage from app.",
        ],
        "Bulk create from Excel on Web only.",
    ),
    (
        "Q-08",
        "Reject requires reason; owner resubmit or delete",
        [
            "PortalAdmin Reject without reason → validation error.",
            "PortalAdmin Reject with reason → Rejected; reason stored and shown to owner.",
            "Owner edits Rejected question and submits → PendingReview.",
            "Alternatively owner deletes Rejected question → allowed.",
        ],
        "Reject is reasoned; owner can fix/resubmit or remove.",
    ),
    (
        "Q-09",
        "Parent / SchoolAdmin ownership consistency",
        [
            "Parent creates → edits own Pending → cannot edit after Approve.",
            "SchoolAdmin same; cannot approve/reject/activate others’ questions.",
        ],
        "Ownership + PortalAdmin-only governances consistent.",
    ),
    (
        "Q-10",
        "Fill in the Blanks — accepted answers + visibility",
        [
            "Creator adds Fill in the Blanks with ≥1 accepted answer (case / partial / length flags as needed).",
            "Create with zero accepted answers → validation error.",
            "PortalAdmin Approves → manage roles see accepted answers in bank detail.",
            "Student starts quiz attempt containing the Fill question → API/UI must NOT return accepted/model answers before submit.",
            "After submit (or teacher/AI review per flags): scoring/review may use accepted answers; optional AI then teacher review.",
        ],
        "Fill uses accepted answers; students never see model answers pre-submit.",
    ),
    (
        "Q-11",
        "Choice-type validation",
        [
            "Single Choice with 0 or 2+ correct options → rejected.",
            "Multiple Choice with no correct options → rejected.",
            "True/False with wrong option count or both correct → rejected.",
        ],
        "Per-type IsCorrect rules enforced on create/update/import.",
    ),
]

CHECKLIST = [
    ("1", "PortalAdmin Create/Update/Delete/Approve/Reject(reason)/Activate/Deactivate/Archive"),
    ("2", "SchoolAdmin / CampusAdmin / Teacher / Parent can Create"),
    ("3", "Those roles Update own only if not Approved (incl. Rejected → resubmit)"),
    ("4", "Those roles Delete own Draft/PendingReview/Rejected only"),
    ("5", "Only PortalAdmin deletes Approved (or any)"),
    ("6", "Only PortalAdmin Approves; Reject requires reason"),
    ("7", "Only PortalAdmin Activate / Deactivate / Archive"),
    ("8", "Manage roles see all Approved + answers; Student has no bank access"),
    ("9", "Student still sees questions inside quizzes"),
    ("10", "Status IDs 110–114; Type 100–104; Difficulty 2001–2003"),
    ("11", "CampusAdmin included in create/manage scope"),
    ("12", "Excel import on Web only (not Mobile)"),
    ("13", "NOW types: Single Choice, Multiple Choice, True/False, Fill in the Blanks"),
    ("14", "Fill uses question_accepted_answers (≥1); choice types use question_options"),
    ("15", "Student attempt: no model/accepted answers before submit (esp. Fill)"),
    ("16", "Descriptive + future types hidden until enabled"),
    ("17", "IsAiApproved is optional flag — not a second quiz-eligibility gate"),
]

OPEN_DECISIONS = [
    "Import default status: Draft (implemented) — owner submits to PendingReview.",
    "Delete while linked to quizzes: blocked for everyone including PortalAdmin (implemented).",
    "Rejected edit stays Rejected until Submit for review (implemented).",
    "IsAiApproved dual gate dropped — PortalAdmin Approve sets quiz-eligible (implemented).",
    "Fill in the Blanks is NOW — store answers in question_accepted_answers (not options). Descriptive remains NOT NOW.",
    "Fill AI/Teacher review: optional attempt-time flags; bank approval stays PortalAdmin-only.",
]

# Product improvement suggestions (advisory — not blockers)
SUGGESTIONS = [
    (
        "Wire Fill to accepted answers",
        "Stop storing Fill answers as question_options. Create/update/import must write QuestionAcceptedAnswer rows; migrate any legacy Fill options into accepted answers.",
    ),
    (
        "Align AllowAIReview / AllowTeacherReview",
        "Product model uses boolean AllowAIReview / AllowTeacherReview on accepted answers. Current entity has AiReview / TeacherReview strings — add bool flags (keep review text on attempt/review tables if needed).",
    ),
    (
        "Enforce type validation in API",
        "Single Choice: exactly one IsCorrect. Multiple Choice: ≥1 IsCorrect. True/False: two options, one correct. Fill: ≥1 accepted answer. Reject invalid create/update/import rows.",
    ),
    (
        "Hide NOT NOW / Future types",
        "Create UI, Mobile, and Excel import should only offer 100–103 until Descriptive is enabled. Do not add Matching/Ordering/etc. lookups until built.",
    ),
    (
        "Attempt answer privacy",
        "For Fill (and later Descriptive): never return AnswerText / NormalizedAnswer / Explanation of correct answers on student attempt GET until after submission (or explicit review stage).",
    ),
    (
        "Canonical status IDs only",
        "Use QuestionStatus IDs 110–114 (Draft, PendingReview, Approved, Rejected, Archived). Writes prefer those IDs; legacy aliases remain readable for old rows.",
    ),
    (
        "Reject reason UX",
        "Require min length (e.g. 10 chars), show reason on owner’s detail card, and clear reason when resubmitted to PendingReview.",
    ),
    (
        "Resubmit as explicit action",
        "Prefer Edit (stays Rejected) + Submit for review → PendingReview, rather than auto-flipping status on every save.",
    ),
    (
        "Import dry-run",
        "Use Dry run Excel before Import; review the full row-error list; Confirm import after a clean dry run.",
    ),
    (
        "Correct-answer visibility",
        "Show options/accepted answers in bank detail for manage roles; Students only via quiz attempt rules after submit.",
    ),
    (
        "Audit trail",
        "Log Approve/Reject/Activate/Archive with actor + timestamp (and reject reason) for disputes and QA.",
    ),
]

REGENERATE = (
    "Generated from docs/03_build_questions_business_qa.py. "
    "Edit that file and re-run to refresh HTML + DOCX. "
    "Fill in the Blanks is in NOW scope; Descriptive remains NOT NOW."
)

# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

CSS = """
    :root {
      --ink: #0b2545; --blue: #2e74b5; --muted: #53606d; --border: #d9e1ea;
      --fill: #f4f6f9; --ok: #166534; --ok-bg: #ecfdf5; --warn: #92400e; --warn-bg: #fffbeb;
    }
    body { margin: 0; color: #17212b; font-family: Calibri, Arial, sans-serif; font-size: 15px; line-height: 1.55; }
    main { max-width: 1100px; margin: 0 auto; padding: 40px 24px 64px; }
    header { border-bottom: 1px solid var(--border); margin-bottom: 28px; padding-bottom: 18px; }
    h1 { margin: 0 0 8px; color: var(--ink); font-size: 30px; }
    .subtitle { margin: 0; color: var(--muted); font-size: 16px; }
    .meta { margin-top: 12px; display: flex; flex-wrap: wrap; gap: 8px; }
    .chip {
      display: inline-block; border-radius: 999px; padding: 4px 10px; font-size: 12px;
      font-weight: 600; background: var(--fill); border: 1px solid var(--border); color: var(--ink);
    }
    .chip.warn { background: var(--warn-bg); color: var(--warn); border-color: #fde68a; }
    h2 { color: var(--blue); margin: 34px 0 10px; font-size: 22px; border-bottom: 1px solid var(--border); padding-bottom: 6px; }
    h3 { margin: 22px 0 8px; color: var(--ink); font-size: 17px; }
    table { width: 100%; border-collapse: collapse; margin: 12px 0 18px; font-size: 13.5px; }
    th, td { border: 1px solid var(--border); padding: 8px 10px; vertical-align: top; text-align: left; }
    th { background: var(--fill); color: var(--ink); }
    code { font-family: Consolas, monospace; font-size: 12.5px; background: #f1f5f9; padding: 1px 5px; border-radius: 4px; }
    .note { margin: 12px 0; padding: 10px 12px; border-radius: 8px; background: var(--warn-bg); border: 1px solid #fde68a; color: var(--warn); }
    .ok { margin: 12px 0; padding: 10px 12px; border-radius: 8px; background: var(--ok-bg); border: 1px solid #bbf7d0; color: var(--ok); }
    .scenario { border: 1px solid var(--border); border-radius: 12px; padding: 14px 16px; margin: 14px 0; }
    .scenario h3 { margin-top: 0; }
    .id { font-weight: 700; color: var(--blue); margin-right: 6px; }
    .expect { margin-top: 10px; padding: 10px 12px; border-radius: 8px; background: var(--ok-bg); border: 1px solid #bbf7d0; color: var(--ok); }
    .toc a { color: var(--blue); text-decoration: none; }
    footer { margin-top: 40px; padding-top: 16px; border-top: 1px solid var(--border); color: var(--muted); font-size: 13px; }
    ul, ol { margin: 8px 0 16px 22px; }
"""


def _e(text: str) -> str:
    return html.escape(str(text), quote=True)


def _table(headers: list[str], rows: list[tuple], widths: list[str] | None = None) -> str:
    ths = []
    for i, h in enumerate(headers):
        style = f' style="width:{widths[i]}"' if widths and i < len(widths) else ""
        ths.append(f"<th{style}>{_e(h)}</th>")
    body = "".join(
        "<tr>" + "".join(f"<td>{_e(c)}</td>" for c in row) + "</tr>" for row in rows
    )
    return f"<table><thead><tr>{''.join(ths)}</tr></thead><tbody>{body}</tbody></table>"


def build_html() -> str:
    scenarios_html = []
    for sid, title, steps, expect in SCENARIOS:
        lis = "".join(f"<li>{_e(s)}</li>" for s in steps)
        scenarios_html.append(
            f'<div class="scenario"><h3><span class="id">{_e(sid)}</span>{_e(title)}</h3>'
            f'<ol class="steps">{lis}</ol>'
            f'<div class="expect">Expected: {_e(expect)}</div></div>'
        )

    rules_lis = "".join(
        f"<li><strong>{_e(n)}. {_e(t)}:</strong> {_e(d)}</li>" for n, t, d in TARGET_RULES
    )
    decisions = "".join(f"<li>{_e(d)}</li>" for d in OPEN_DECISIONS)
    suggestions_html = "".join(
        f"<li><strong>{_e(t)}:</strong> {_e(d)}</li>" for t, d in SUGGESTIONS
    )
    checklist = _table(
        ["#", "Check", "Pass?"],
        [(a, b, "☐") for a, b in CHECKLIST],
        ["8%", "82%", "10%"],
    )

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{_e(META['title'])}</title>
  <style>{CSS}</style>
</head>
<body>
<main>
  <header>
    <h1>{_e(META['title'])}</h1>
    <p class="subtitle">{_e(META['subtitle'])}</p>
    <div class="meta">
      <span class="chip">{_e(META['prepared_for'])}</span>
      <span class="chip">{_e(META['prepared_on'])}</span>
      <span class="chip warn">{_e(META['status'])}</span>
      <span class="chip">{_e(META['roles'])}</span>
    </div>
  </header>

  <h2 id="toc">Table of contents</h2>
  <ol class="toc">
    <li><a href="#target">1. Target business rules</a></li>
    <li><a href="#matrix">2. Permissions matrix</a></li>
    <li><a href="#lookups">3. Lookup IDs (QuestionStatus / Type / Difficulty)</a></li>
    <li><a href="#types">4. Question types — Now / Not Now / Future</a></li>
    <li><a href="#model">5. Question data model</a></li>
    <li><a href="#lifecycle">6. Status lifecycle</a></li>
    <li><a href="#import">7. Excel import (Web only)</a></li>
    <li><a href="#gap">8. Gap vs current codebase</a></li>
    <li><a href="#decisions">9. Open decisions</a></li>
    <li><a href="#suggestions">10. Suggested improvements</a></li>
    <li><a href="#scenarios">11. QA scenarios</a></li>
    <li><a href="#checklist">12. Verification checklist</a></li>
  </ol>

  <div class="ok">
    <strong>Permissions implemented</strong> in WebApi + React Web: PortalAdmin-only approve/reject/lifecycle,
    CampusAdmin create scope, owner pre-approval edit/delete, Reject reason required + Submit for review,
    Excel import (Web), quiz eligibility after single Approve. Mobile has no Excel import.
  </div>
  <div class="note">
    <strong>Type scope (15 Jul 2026):</strong> NOW = Single Choice, Multiple Choice, True/False,
    and <strong>Fill in the Blanks</strong> (accepted answers). Descriptive = NOT NOW.
    Matching / Ordering / media / math / coding / case-study = Future.
  </div>

  <h2 id="target">1. Target business rules</h2>
  <p><strong>Admin Portal</strong> in product language = role <code>PortalAdmin</code> (lookup id 2010).</p>
  <ol>{rules_lis}</ol>
  <p class="ok">Rule numbers match the product request (item 6 was unused in the source list).</p>

  <h2 id="matrix">2. Permissions matrix (target)</h2>
  {_table(PERMISSION_HEADERS, PERMISSION_ROWS)}
  <ul>
    <li><strong>Student:</strong> no question-bank access. Questions appear only inside quiz attempts.</li>
    <li><strong>Reject:</strong> PortalAdmin must enter a non-empty rejection reason.</li>
    <li><strong>After Reject:</strong> owner may edit + resubmit to <code>PendingReview</code>, or delete.</li>
    <li><strong>Excel import:</strong> Web only — not on Mobile.</li>
  </ul>

  <h2 id="lookups">3. Lookup IDs</h2>
  <h3>QuestionStatus</h3>
  {_table(["ID", "Name", "Type", "Notes"], STATUS_ROWS)}
  <h3>QuestionType</h3>
  {_table(["ID", "Name", "Type", "Sort", "Scope"], TYPE_ROWS)}
  <h3>DifficultyLevel</h3>
  {_table(["ID", "Name", "Type", "Sort"], DIFFICULTY_ROWS)}

  <h2 id="types">4. Question types — Now / Not Now / Future</h2>
  <p>Quiz Module type support and answer storage:</p>
  {_table(["Question type", "Scope", "Answer storage", "Validation / rules"], TYPE_RULE_ROWS)}
  <h3>Future types (roadmap — no create UI yet)</h3>
  {_table(["ID", "Name", "Scope", "Notes"], FUTURE_TYPE_ROWS)}
  <ul>
    <li><strong>Fill in the Blanks (NOW):</strong> AI may provide an initial attempt review; teacher review may be required;
      model / accepted answers must not be shown before the student submits.</li>
    <li><strong>Descriptive (NOT NOW):</strong> same review/visibility pattern when implemented; hide from create/import until then.</li>
    <li><strong>Bank approval</strong> remains PortalAdmin-only for all types — separate from attempt AI/teacher review.</li>
  </ul>

  <h2 id="model">5. Question data model</h2>
  <h3>Table: questions</h3>
  {_table(["Field", "Notes"], QUESTION_FIELD_ROWS)}
  <h3>Table: question_options</h3>
  <p>For choice-based types (Single Choice, Multiple Choice, True/False):</p>
  {_table(["Field", "Notes"], OPTION_FIELD_ROWS)}
  <h3>Table: question_accepted_answers</h3>
  <p>For Fill in the Blanks (NOW) and Descriptive (NOT NOW):</p>
  {_table(["Field", "Notes"], ACCEPTED_ANSWER_FIELD_ROWS)}
  <div class="note">
    Do not store Fill answers as <code>question_options</code>. Use <code>question_accepted_answers</code> only.
    Manage roles may see accepted answers in the bank; students must not see them before attempt submit.
  </div>

  <h2 id="lifecycle">6. Status lifecycle (target)</h2>
  {_table(["Step", "Status", "Notes"], LIFECYCLE)}
  <ul>
    <li>Canonical status names/IDs above replace loose aliases where possible.</li>
    <li>Approved questions show options / accepted answers to manage roles (not Students in the bank).</li>
  </ul>

  <h2 id="import">7. Excel import (Web only)</h2>
  <p>On <strong>Web</strong>, create-capable roles get an <strong>Import from Excel</strong> action.
  <strong>Mobile does not</strong> expose import.</p>
  {_table(["Column", "Required", "Notes"], IMPORT_COLUMNS)}
  <div class="note">Import must never set status to Approved. Only PortalAdmin approval promotes to Approved.
  Import should accept NOW types only (100–103); reject Descriptive until enabled.</div>

  <h2 id="gap">8. Gap vs current codebase</h2>
  <p>Core permission / approve / import rules are implemented. Type/model work for Fill accepted answers
  and stricter validation remains. Remaining UX polish and Mobile gating are separate.</p>
  {_table(["Topic", "Current", "Target", "Action"], GAP_ROWS)}

  <h2 id="decisions">9. Open decisions</h2>
  <ol>{decisions}</ol>

  <h2 id="suggestions">10. Suggested improvements</h2>
  <p>Advisory recommendations to strengthen question logic:</p>
  <ul>
{suggestions_html}
  </ul>

  <h2 id="scenarios">11. QA scenarios</h2>
  {''.join(scenarios_html)}

  <h2 id="checklist">12. Verification checklist</h2>
  {checklist}

  <footer>{_e(REGENERATE)}</footer>
</main>
</body>
</html>
"""


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, bottom=60, start=80, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_table_borders(table, color="D9E1EA"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "F4F6F9")
        p = table.rows[0].cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(11, 37, 69)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_note(doc, label, text, fill="FFFBEB"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, "FDE68A")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    doc.add_paragraph()


def set_styles(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for name, size, color in (
        ("Title", 26, RGBColor(11, 37, 69)),
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color


def build_docx() -> None:
    doc = Document()
    set_styles(doc)
    doc.add_paragraph(META["title"], style="Title")
    doc.add_paragraph(META["subtitle"])
    add_table(
        doc,
        ["Prepared for", "Date", "Status"],
        [[META["prepared_for"], META["prepared_on"], META["status"]]],
        [2800, 1800, 4760],
    )
    add_note(
        doc,
        "Permissions implemented",
        "WebApi + React Web: PortalAdmin-only approve/reject/lifecycle, CampusAdmin create scope, owner pre-approval edit/delete, Reject reason required + Submit for review, Excel import (Web), quiz eligibility after single Approve. Mobile has no Excel import.",
        "EDF8F2",
    )
    add_note(
        doc,
        "Type scope (15 Jul 2026)",
        "NOW = Single Choice, Multiple Choice, True/False, and Fill in the Blanks (accepted answers). Descriptive = NOT NOW. Matching / Ordering / media / math / coding / case-study = Future.",
        "FFF7ED",
    )

    doc.add_heading("1. Target business rules", level=1)
    doc.add_paragraph("Admin Portal = PortalAdmin (lookup id 2010).")
    for n, t, d in TARGET_RULES:
        add_number(doc, f"{n}. {t}: {d}")

    doc.add_heading("2. Permissions matrix (target)", level=1)
    add_table(
        doc,
        PERMISSION_HEADERS,
        PERMISSION_ROWS,
        [1800, 1200, 1100, 1100, 1000, 1000, 2160],
    )
    add_bullet(doc, "Student: no question-bank access — questions only inside quizzes.")
    add_bullet(doc, "Reject: PortalAdmin must enter a non-empty rejection reason.")
    add_bullet(doc, "After Reject: owner may edit + resubmit to PendingReview, or delete.")
    add_bullet(doc, "Excel import: Web only — not on Mobile.")

    doc.add_heading("3. Lookup IDs", level=1)
    doc.add_heading("QuestionStatus", level=2)
    add_table(doc, ["ID", "Name", "Type", "Notes"], STATUS_ROWS, [900, 1600, 1800, 5060])
    doc.add_heading("QuestionType", level=2)
    add_table(
        doc,
        ["ID", "Name", "Type", "Sort", "Scope"],
        TYPE_ROWS,
        [700, 2000, 1600, 700, 4360],
    )
    doc.add_heading("DifficultyLevel", level=2)
    add_table(doc, ["ID", "Name", "Type", "Sort"], DIFFICULTY_ROWS, [900, 1800, 2200, 4460])

    doc.add_heading("4. Question types — Now / Not Now / Future", level=1)
    doc.add_paragraph("Quiz Module type support and answer storage:")
    add_table(
        doc,
        ["Question type", "Scope", "Answer storage", "Validation / rules"],
        TYPE_RULE_ROWS,
        [1800, 1000, 1800, 4760],
    )
    doc.add_heading("Future types (roadmap — no create UI yet)", level=2)
    add_table(doc, ["ID", "Name", "Scope", "Notes"], FUTURE_TYPE_ROWS, [700, 2200, 1200, 5260])
    add_bullet(
        doc,
        "Fill in the Blanks (NOW): AI may provide an initial attempt review; teacher review may be required; model/accepted answers must not be shown before the student submits.",
    )
    add_bullet(
        doc,
        "Descriptive (NOT NOW): same review/visibility pattern when implemented; hide from create/import until then.",
    )
    add_bullet(
        doc,
        "Bank approval remains PortalAdmin-only for all types — separate from attempt AI/teacher review.",
    )

    doc.add_heading("5. Question data model", level=1)
    doc.add_heading("Table: questions", level=2)
    add_table(doc, ["Field", "Notes"], QUESTION_FIELD_ROWS, [2800, 6560])
    doc.add_heading("Table: question_options", level=2)
    doc.add_paragraph("For choice-based types (Single Choice, Multiple Choice, True/False):")
    add_table(doc, ["Field", "Notes"], OPTION_FIELD_ROWS, [2800, 6560])
    doc.add_heading("Table: question_accepted_answers", level=2)
    doc.add_paragraph("For Fill in the Blanks (NOW) and Descriptive (NOT NOW):")
    add_table(doc, ["Field", "Notes"], ACCEPTED_ANSWER_FIELD_ROWS, [2800, 6560])
    add_note(
        doc,
        "Important",
        "Do not store Fill answers as question_options. Use question_accepted_answers only. Manage roles may see accepted answers in the bank; students must not see them before attempt submit.",
    )

    doc.add_heading("6. Status lifecycle (target)", level=1)
    add_table(doc, ["Step", "Status", "Notes"], LIFECYCLE, [2200, 1800, 5360])

    doc.add_heading("7. Excel import (Web only)", level=1)
    doc.add_paragraph(
        "On Web, create-capable roles get Import from Excel. Mobile does not expose import."
    )
    add_table(doc, ["Column", "Required", "Notes"], IMPORT_COLUMNS, [2000, 1400, 5960])
    add_note(
        doc,
        "Important",
        "Import must never set status to Approved. Import should accept NOW types only (100–103); reject Descriptive until enabled.",
    )

    doc.add_heading("8. Gap vs current codebase", level=1)
    add_table(doc, ["Topic", "Current", "Target", "Action"], GAP_ROWS, [1600, 2600, 3200, 1960])

    doc.add_heading("9. Open decisions", level=1)
    for d in OPEN_DECISIONS:
        add_bullet(doc, d)

    doc.add_heading("10. Suggested improvements", level=1)
    doc.add_paragraph("Advisory recommendations to strengthen question logic:")
    for title, detail in SUGGESTIONS:
        add_bullet(doc, f"{title}: {detail}")

    doc.add_heading("11. QA scenarios", level=1)
    for sid, title, steps, expect in SCENARIOS:
        doc.add_heading(f"{sid} — {title}", level=2)
        for s in steps:
            add_number(doc, s)
        p = doc.add_paragraph()
        r = p.add_run("Expected: ")
        r.bold = True
        p.add_run(expect)

    doc.add_heading("12. Verification checklist", level=1)
    add_table(
        doc,
        ["#", "Check", "Pass?"],
        [(a, b, "☐") for a, b in CHECKLIST],
        [600, 7600, 1160],
    )

    doc.add_paragraph(REGENERATE)
    doc.core_properties.title = META["title"]
    doc.core_properties.subject = "Questions business rules, type scope, and QA"
    doc.core_properties.comments = "Fill in the Blanks is NOW; Descriptive NOT NOW."
    doc.save(DOCX_OUTPUT)


def main() -> None:
    HTML_OUTPUT.write_text(build_html(), encoding="utf-8")
    build_docx()
    print(f"Wrote {HTML_OUTPUT}")
    print(f"Wrote {DOCX_OUTPUT}")


if __name__ == "__main__":
    main()
