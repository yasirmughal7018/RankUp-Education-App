"""Generate RankUp Education Authentication & Login Logic DOCX documentation.

Source of truth aligned with AuthService + React (Aug 2026).
Also maintain docs/02_RankUp_Authentication_Logic.html in sync when editing.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parent
OUTPUT = DOCS_DIR / "02_RankUp_Authentication_Logic.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_table_borders(table, color="B7C6D8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
        p = table.rows[0].cells[idx].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(22, 57, 87)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].paragraphs[0].add_run(value)
    doc.add_paragraph()


def add_note(doc, label, text, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    doc.add_paragraph()


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, "2F3A4A")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "172033")
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(248, 250, 252)
    doc.add_paragraph()


def set_document_styles(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor(23, 50, 77)
    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(91, 100, 114)
    for level, size in ((1, 16), (2, 13), (3, 12)):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(46, 116, 181)
        h.font.bold = True


def build_doc():
    doc = Document()
    set_document_styles(doc)

    doc.add_paragraph("RankUp Education — Authentication & Login Logic", style="Title")
    doc.add_paragraph(
        "Source of truth for the shared Web API auth model, with explicit "
        "React vs Flutter Mobile client coverage.",
        style="Subtitle",
    )

    add_table(
        doc,
        ["Prepared for", "Scope", "Auth model"],
        [[
            "RankUp Education workspace",
            "API shared; React full; Flutter partial (§12)",
            "JWT access + hashed refresh; multi-role",
        ]],
        [2200, 3200, 3960],
    )

    add_note(
        doc,
        "Supersedes",
        "Prior drafts that used admin_target, single app_users.role, SuperAdmin naming, "
        "admin-set password on approve, and hard-delete reject are not the current model. "
        "Companion QA: docs/03_RankUp_User_Creation_Approval_QA.html.",
        "EEF8F1",
    )

    add_note(
        doc,
        "Client parity",
        "API contract is shared. React = full admin/web (directory, school-change apply queue, "
        "password-reset clear, email reset-password page). Flutter Mobile = login-status (incl. lock), "
        "set-initial-password, request access, password-reset request, school-change request, logout revoke — "
        "not Web-only admin surfaces. Do not assume Flutter matches every React screen.",
        "FFF8E8",
    )

    # --- 1 ---
    doc.add_heading("1. Executive Summary", level=1)
    for item in [
        "Authenticated API calls use a Bearer JWT access token (~30 min). Refresh tokens last 30 days and are stored hashed.",
        "Two-step login: POST /api/auth/login-status (identifier) → set-initial-password or password login.",
        "Username = email for Student / Parent / Teacher / SchoolAdmin / CampusAdmin (email required on register and directory create). Login lookup: username → cnic → mobile_number.",
        "Self-service registration: inactive app_users + app_user_roles + app_approval. Scoped SchoolAdmin / CampusAdmin can activate Student/Teacher in scope; PortalAdmin can activate any; no-school / Parent → PortalAdmin only. User sets password via set-initial-password.",
        "Forgot password: enter username (email) → email reset link + pending app_user_password_reset_request + notify role-scoped helpers. First completion wins (email complete or admin/parent clear).",
        "Reject is soft (rejected_at + required rejection_reason). Unique indexes ignore rejected rows so the person can re-request.",
        "Roles live in app_user_roles (multi-role where allowed). Session role on JWT and refresh_tokens.active_role "
        "(UI “Current / Acting as” — not an is_active column on role rows). Student and PortalAdmin are exclusive; "
        "SchoolAdmin/CampusAdmin/Teacher/Parent may combine (typical: Parent+Teacher).",
        "Users may request Parent/Teacher as an extra role (app_user_role_request; account stays active) or remove "
        "Parent/Teacher when another role remains. Directory admins may grant Parent↔Teacher directly.",
        "School/campus change (Teacher/Student/CampusAdmin) locks the account until destination apply or reject. Parent cannot request school/campus change. Student reject may optionally leaveWithoutSchool.",
        "admin_target and app_users.role columns are dropped / unused.",
        "Six roles: PortalAdmin, SchoolAdmin, CampusAdmin, Teacher, Student, Parent.",
    ]:
        add_bullet(doc, item)

    # --- 2 ---
    doc.add_heading("2. Username & Login Rules", level=1)
    doc.add_heading("2.1 Username = email", level=2)
    for item in [
        "Self-service register (Student / Parent / Teacher): email is required; username = normalized email.",
        "Directory create (Student / Teacher / Parent / SchoolAdmin / CampusAdmin): email is required; username = email.",
        "Username stays as registered email on activation (no rewrite to CNIC).",
        "Mobile and CNIC remain optional contact / alternate login identifiers (mobile not required on register).",
        "Student roll number is required only when a school is selected.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("2.2 Login identifiers", level=2)
    add_code(
        doc,
        "GetByLoginIdentifierAsync(identifier):\n"
        "  1. Match app_users.username (usually email)\n"
        "  2. Else match app_users.cnic\n"
        "  3. Else match app_users.mobile_number\n"
        "Prefer non-rejected rows (PreferActiveRegistrationAsync), then newest Id.\n"
        "Clients send the identifier in the username JSON field.",
    )

    doc.add_heading("2.3 Uniqueness (soft-reject aware)", level=2)
    add_table(
        doc,
        ["Check / index", "Rule"],
        [
            ["UsernameExists / ix_app_users_username_active", "Unique where rejected_at IS NULL"],
            ["CnicExists / ix_app_users_cnic_active", "Unique among non-rejected when CNIC set"],
            ["MobileNumberExistsAsync", "Exists check skips rejected rows"],
        ],
        [4200, 5160],
    )

    doc.add_heading("2.4 Parent vs Student / Teacher routing", level=2)
    add_table(
        doc,
        ["Role / path", "School / campus", "Approval queue at submit"],
        [
            ["Parent", "None", "PortalAdmin only"],
            ["Student/Teacher — no school", "Omitted", "PortalAdmin only"],
            ["Student/Teacher — school only", "School set", "SchoolAdmin + PortalAdmin"],
            ["Student/Teacher — school + campus", "Both set", "CampusAdmin + SchoolAdmin + PortalAdmin"],
        ],
        [2800, 2200, 4360],
    )
    add_bullet(doc, "Routing is written into app_approval at submit — not via admin_target.")

    doc.add_heading("2.5 must_change_password", level=2)
    add_table(
        doc,
        ["Value", "Meaning"],
        [
            ["null or false", "No forced password change"],
            ["true", "Set on activation / directory provision"],
            ["After set-initial-password or change-password", "Cleared to false"],
        ],
        [3600, 5760],
    )
    add_note(
        doc,
        "Note",
        "After set-initial-password the user is Ready — no extra forced change-password modal for that path. "
        "React must-change modal applies when an existing session has mustChangePassword === true.",
    )

    # --- 3 ---
    doc.add_heading("3. Account States & login-status", level=1)
    add_table(
        doc,
        ["State", "DB condition", "login-status", "Can sign in?"],
        [
            [
                "Pending registration",
                "is_active=false AND password NULL AND rejected_at NULL",
                "PendingApproval",
                "No",
            ],
            [
                "Rejected (soft)",
                "Row kept; rejected_at set",
                "Rejected",
                "No — may re-request",
            ],
            [
                "Needs password",
                "is_active=true AND password NULL",
                "NeedsPasswordSetup",
                "No — set password first",
            ],
            [
                "Ready",
                "is_active=true AND password set",
                "Ready",
                "Yes",
            ],
            [
                "Locked (school change)",
                "is_active=false, password set, pending school-change request",
                "LockedPendingSchoolChange",
                "No",
            ],
            [
                "Deactivated",
                "is_active=false, password set, no pending school change",
                "(throws not a status)",
                "No",
            ],
        ],
        [2000, 3200, 2400, 1760],
    )
    add_code(
        doc,
        "IsPendingRegistration = !IsActive && no password && !IsRejectedRegistration\n"
        "NeedsPasswordSetup    = IsActive && no password\n"
        "IsRejectedRegistration = RejectedAt != null",
    )
    add_note(
        doc,
        "Important",
        "Activation is not PortalAdmin-only: SchoolAdmin / CampusAdmin can activate Student/Teacher "
        "in their school/campus scope. No-school and Parent requests remain PortalAdmin-only. "
        "login-status stays PendingApproval until an authorized admin activates.",
        "FFF8E8",
    )

    # --- 4 ---
    doc.add_heading("4. Login, Set Initial Password, Switch Role", level=1)
    doc.add_heading("4.1 Login flow", level=2)
    add_code(
        doc,
        "1. Enter username (email) — or CNIC / mobile if those match\n"
        "2. POST /api/auth/login-status { username }\n"
        "3. Branch: PendingApproval | Rejected | NeedsPasswordSetup |\n"
        "   LockedPendingSchoolChange | Ready\n"
        "   Locked: React /account-locked; Mobile in-login locked step\n"
        "4a. NeedsPasswordSetup → POST set-initial-password → then login\n"
        "4b. Ready → POST /api/auth/login → JWT + refresh; last_login_at updated",
    )

    doc.add_heading("4.2 Password login gates", level=2)
    for item in [
        "GetByLoginIdentifierAsync",
        "EnsureCanLogin() — blocks deleted, rejected, pending, inactive",
        "If NeedsPasswordSetup → error: set password on login screen first",
        "Verify password (PBKDF2)",
        "Issue refresh for active role + access JWT; update last_login_at",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4.3 Set initial password", level=2)
    for item in [
        "POST /api/auth/set-initial-password — anonymous; identifier + new password (min 6).",
        "Requires active user with NeedsPasswordSetup.",
        "Sets password_hash and clears must_change_password to false.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4.4 Multi-role / switch-role / remove-role", level=2)
    add_table(
        doc,
        ["Rule", "Detail"],
        [
            ["Exclusive", "Student — cannot combine with any other role"],
            ["Exclusive", "PortalAdmin — cannot combine with any other role"],
            ["May combine", "SchoolAdmin, CampusAdmin, Teacher, Parent (any mix)"],
            ["Self-remove", "Only Parent or Teacher; account must keep at least one other role"],
        ],
        [2000, 7360],
    )
    add_bullet(doc, "Examples allowed: Teacher+Parent; CampusAdmin+Teacher; SchoolAdmin+CampusAdmin+Teacher/Parent.")
    for item in [
        "Roles stored in app_user_roles only (user_id, role, created_at). No is_active flag on role rows.",
        "Default login active role = earliest assignment in app_user_roles.",
        "POST /api/auth/switch-role (JWT) issues new access + refresh for selected role.",
        "DELETE /api/auth/me/roles/{role} removes Parent/Teacher and re-issues tokens; blocked if student groups reference that user+role.",
        "refresh_tokens.active_role scopes refresh to that role.",
        "Authorization uses the active session role only (CurrentUser.role); CurrentUser.roles = all assignments.",
        "UI: React user-menu Acting as segmented toggle; Account shows Current + Switch/Remove. Mobile Profile switcher when multi-role.",
    ]:
        add_bullet(doc, item)

    # --- 5 ---
    doc.add_heading("5. Logout & Token Revoke", level=1)
    for item in [
        "POST /api/auth/logout — revoke refresh token when provided; missing token is no-op success.",
        "React AuthProvider and Mobile ApiAuthRepository logout send refresh token then clear local session.",
        "Also revoke-all refresh tokens on school-change lock and account deactivate.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("5.1 Forgot password (email + admin/parent clear)", level=2)
    for item in [
        "POST /api/auth/password-reset/request { username } — always success (no existence leak).",
        "If eligible active account with password: cancel prior pending resets; create app_user_password_reset_request; "
        "email reset link when email is known; notify role-scoped helpers (PasswordResetRequest).",
        "POST /api/auth/password-reset/complete { token, newPassword } — user sets password via emailed link.",
        "POST /api/auth/password-reset/clear { username } — admin or linked Parent clears hash → NeedsPasswordSetup (Web UI).",
        "First completion wins: completing the pending request (email or clear) blocks all further resets for that request.",
    ]:
        add_bullet(doc, item)
    add_table(
        doc,
        ["Requester role", "Notified helpers"],
        [
            ["Student", "SchoolAdmin + CampusAdmin (scope) + linked Parents + PortalAdmin"],
            ["Teacher", "SchoolAdmin + CampusAdmin (scope) + PortalAdmin"],
            ["CampusAdmin", "SchoolAdmin (school) + PortalAdmin"],
            ["SchoolAdmin", "PortalAdmin only"],
            ["Parent", "PortalAdmin only"],
        ],
        [2400, 6960],
    )
    add_table(
        doc,
        ["Clearer", "May clear"],
        [
            ["PortalAdmin", "Any user"],
            ["SchoolAdmin", "Student / Teacher / CampusAdmin / Parent in own school (not SchoolAdmin / PortalAdmin)"],
            ["CampusAdmin", "Student / Teacher in own campus"],
            ["Parent", "Linked Student only"],
        ],
        [2400, 6960],
    )
    add_bullet(
        doc,
        "Web: /forgot-password (request) + /reset-password?token=… (email complete). Clear via notification bell. "
        "Mobile: request only (no clear / no email-complete page yet).",
    )

    # --- 6 ---
    doc.add_heading("6. Registration & Pending Detection", level=1)
    add_code(
        doc,
        "Client → POST /api/auth/register\n"
        "  → username = email (required)\n"
        "  → uniqueness among rejected_at IS NULL\n"
        "  → INSERT app_users (inactive, no password)\n"
        "  → INSERT app_user_roles\n"
        "  → INSERT app_approval for eligible reviewers\n"
        "  → notify admins (RegistrationRequest)",
    )
    for item in [
        "Domain pending: User.IsPendingRegistration.",
        "Admin list: GET /api/auth/registrations/pending (PortalAdmin/SchoolAdmin/CampusAdmin) with pendingApprovers.",
        "Login: login-status → PendingApproval.",
        "Directory-created users: active + NeedsPasswordSetup; do not use app_approval.",
    ]:
        add_bullet(doc, item)

    # --- 7 ---
    doc.add_heading("7. Approval Queue & Soft Reject", level=1)
    add_table(
        doc,
        ["Action", "PortalAdmin", "SchoolAdmin", "CampusAdmin"],
        [
            ["List pending", "Yes", "Yes (scoped)", "Yes (scoped)"],
            ["Approve + activate (in scope)", "Any", "Student/Teacher in school", "Student/Teacher in campus"],
            ["Activate no-school / Parent", "Yes only", "No", "No"],
            ["Soft-reject", "Yes", "Yes (scoped)", "Yes (scoped)"],
        ],
        [2800, 2000, 2280, 2280],
    )

    doc.add_heading("7.2 Activation hierarchy", level=2)
    for item in [
        "PortalAdmin: activate any pending registration (including Parent / no-school).",
        "SchoolAdmin: activate Student/Teacher for their school (school-only or with campus).",
        "CampusAdmin: activate Student/Teacher for their campus.",
        "No school on the request (and Parent): PortalAdmin only — School/Campus cannot review.",
        "Approve body is empty — no admin password payload. Username remains the registered email.",
        "UI may show “Approved — awaiting activation” only if a reviewer recorded approval without activating (rare for in-scope School/Campus Student/Teacher).",
    ]:
        add_bullet(doc, item)

    doc.add_heading("7.3 Pending with filter", level=2)
    for item in [
        "pendingApprovers = queue rows where is_approved IS NULL (ListPendingApproversForUserAsync).",
        "If any CampusAdmin has already approved, pending SchoolAdmin rows may be omitted from “Pending with” display.",
        "Once any authorized admin activates, the request leaves the pending list.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("7.4 Soft reject", level=2)
    for item in [
        "Body { reason } required (min 10 chars).",
        "Soft reject: set rejected_at + rejection_reason; keep row + approval trail; login-status Rejected (may show reason); user may submit a new request.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("7b. Additional role requests & remove role", level=1)
    for item in [
        "POST /api/auth/me/role-requests — request Parent or Teacher as additional role (does not lock account).",
        "Teacher request needs school/campus (+ optional teacher code). Creates app_user_role_request; notifies RoleRequest.",
        "GET /api/auth/me may include pendingRoleRequest while open.",
        "Admin: GET /api/auth/role-requests/pending; POST …/approve; POST …/reject { reason }. Web: /admin/directory/role-requests.",
        "Directory grant: POST /api/directory/parents/{id}/roles/teacher; POST /api/directory/teachers/{id}/roles/parent.",
        "DELETE /api/auth/me/roles/{role} — self-remove Parent/Teacher when another role remains (Web Account Remove).",
    ]:
        add_bullet(doc, item)

    # --- 8 ---
    doc.add_heading("8. School / Campus Change Lock", level=1)
    for item in [
        "Who may request: Teacher, Student, CampusAdmin. Cannot: PortalAdmin, SchoolAdmin, Parent.",
        "POST /api/auth/me/school-change → pending request → is_active=false → revoke all refresh tokens → locked messaging.",
        "login-status → LockedPendingSchoolChange.",
        "Request UI: React /account and Flutter Mobile Profile (eligible roles).",
        "Admin apply/reject queue: Web only — /admin/directory/school-changes. Mobile does not implement it.",
    ]:
        add_bullet(doc, item)

    add_table(
        doc,
        ["Approver", "Can apply (unlock + move)"],
        [
            ["PortalAdmin", "Any pending request"],
            ["SchoolAdmin", "Teacher / Student / CampusAdmin into destination school"],
            ["CampusAdmin", "Teacher / Student into own campus (not CampusAdmin peers)"],
        ],
        [2400, 6960],
    )
    add_bullet(
        doc,
        "List filters match this matrix: every admin who can see a pending request can Approve & apply (or Reject). "
        "Reject unlocks without changing school/campus (optional leaveWithoutSchool for Student clears school/campus). "
        "Admin UI (Web): /admin/directory/school-changes.",
    )

    # --- 9 ---
    doc.add_heading("9. Supporting Auth Tables", level=1)
    doc.add_heading("9.1 app_users (auth-relevant)", level=2)
    add_table(
        doc,
        ["Column", "Purpose"],
        [
            ["username", "Email (login username)"],
            ["email", "Required contact; same as username for new accounts"],
            ["mobile_number / cnic", "Optional contact + alternate login ids"],
            ["password_hash", "NULL until set-initial-password / email complete"],
            ["is_active", "Pending / locked / deactivated gate"],
            ["must_change_password", "Force / setup signalling"],
            ["rejected_at", "Soft-reject; NULL = not rejected"],
            ["rejection_reason", "Required text when soft-rejected"],
            ["requested_at", "Registration submit time"],
            ["school_id / campus_id", "Scope owned on user"],
            ["roll_number_teacher_code", "Student roll or teacher code"],
            ["avatar_url / last_login_at", "Profile image; last successful login"],
        ],
        [3200, 6160],
    )
    add_bullet(doc, "Removed: role and admin_target on app_users.")

    doc.add_heading("9.2 Role, approval, school-change, password-reset", level=2)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["app_user_roles", "Multi-role assignments (source of truth; no is_active)"],
            ["app_user_role_request", "Pending self-service Parent/Teacher role requests"],
            ["app_approval", "Registration approval queue per reviewer"],
            ["app_user_school_change_request", "Pending moves + lock"],
            ["app_user_school_change_approval", "Reviewer trail for school-change"],
            ["app_user_password_reset_request", "Forgot-password pending request; first completion wins"],
        ],
        [3600, 5760],
    )

    doc.add_heading("9.3 Session & notify", level=2)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["refresh_tokens", "Hashed refresh tokens; includes active_role"],
            ["device_sessions", "Device session tracking"],
            ["app_notifications", "In-app alerts (e.g. RegistrationRequest)"],
        ],
        [2800, 6560],
    )

    doc.add_heading("9.4 Slim profiles", level=2)
    add_table(
        doc,
        ["Table", "Keeps", "On app_users"],
        [
            ["app_user_students", "grade, section, mobile", "school, campus, cnic, roll"],
            ["app_user_teachers", "mobile", "school, campus, cnic, teacher code"],
            ["app_user_parents", "mobile", "cnic"],
        ],
        [2400, 2800, 4160],
    )

    # --- 10 ---
    doc.add_heading("10. API Endpoints", level=1)
    add_table(
        doc,
        ["Method", "Route", "Auth", "Description"],
        [
            ["POST", "/api/auth/login-status", "Anonymous", "Identifier → status machine"],
            ["POST", "/api/auth/set-initial-password", "Anonymous", "First password after approval"],
            ["POST", "/api/auth/login", "Anonymous", "Email/CNIC/mobile + password → tokens"],
            ["POST", "/api/auth/switch-role", "JWT", "Tokens for another assigned role"],
            ["DELETE", "/api/auth/me/roles/{role}", "JWT", "Remove Parent/Teacher; re-issue tokens"],
            ["POST", "/api/auth/me/role-requests", "JWT", "Request additional Parent or Teacher"],
            ["GET", "/api/auth/role-requests/pending", "Admin", "List pending additional-role requests"],
            ["POST", "/api/auth/role-requests/{id}/approve", "Admin", "Approve additional role"],
            ["POST", "/api/auth/role-requests/{id}/reject", "Admin", "Reject additional role (reason)"],
            ["POST", "/api/auth/register", "Anonymous", "Pending user + approval queue (email username)"],
            ["GET", "/api/auth/registration-options/schools", "Anonymous", "School dropdown"],
            ["GET", "/api/auth/registration-options/schools/{id}/campuses", "Anonymous", "Campus dropdown"],
            ["GET", "/api/auth/registrations/pending", "Admin", "List pending (scoped)"],
            ["POST", "/api/auth/registrations/{id}/approve", "Admin", "Approve; activate when authorized"],
            ["POST", "/api/auth/registrations/{id}/reject", "Admin", "Soft-reject with required reason"],
            ["POST", "/api/auth/change-password", "JWT", "Change password; clear must-change"],
            ["POST", "/api/auth/token/refresh", "Anonymous", "Rotate access via refresh"],
            ["POST", "/api/auth/logout", "Anonymous", "Revoke refresh if provided"],
            ["GET", "/api/auth/me", "JWT", "Current user profile"],
            ["PUT", "/api/auth/me", "JWT", "Update profile"],
            ["POST", "/api/auth/me/school-change", "JWT", "Request move + lock"],
            ["POST", "/api/auth/me/avatar", "JWT", "Upload avatar"],
            ["POST", "/api/auth/me/deactivate", "JWT", "Self-deactivate"],
            ["GET", "/api/auth/school-changes/pending", "Admin", "Pending queue (Web UI)"],
            ["POST", "/api/auth/school-changes/{id}/approve", "Admin", "Apply unlock (Web UI)"],
            ["POST", "/api/auth/school-changes/{id}/reject", "Admin", "Reject + unlock; optional leaveWithoutSchool (Student)"],
            ["POST", "/api/auth/password-reset/request", "Anonymous", "Email link + notify helpers (no existence leak)"],
            ["POST", "/api/auth/password-reset/complete", "Anonymous", "Set password via emailed token (first wins)"],
            ["POST", "/api/auth/password-reset/clear", "Admin/Parent", "Clear pending reset → NeedsPasswordSetup (Web UI)"],
        ],
        [900, 3800, 1400, 3260],
    )

    # --- 11 ---
    doc.add_heading("11. JWT, Refresh Tokens, Security", level=1)
    add_table(
        doc,
        ["Claim", "Source"],
        [
            ["sub / userId", "app_users.id"],
            ["name", "app_users.username"],
            ["role", "Active session role (enum name)"],
            ["roles", "All assigned roles, comma-separated (app_user_roles)"],
            ["permissions", "AuthPermissions.ForRole(activeRole), comma-separated"],
            ["profileId", "Role profile id when loaded"],
            ["schoolId / campusId", "User school/campus"],
        ],
        [2800, 6560],
    )
    add_bullet(
        doc,
        "role = active session role; roles = full assignment list. Authorization uses active role. "
        "mustChangePassword is on login/me responses, not a JWT claim.",
    )
    for item in [
        "Passwords: PBKDF2 via PasswordHasher.",
        "Refresh tokens stored as SHA-256 hex; lifetime 30 days.",
        "Access token: Jwt:AccessTokenMinutes (default 30).",
        "Login rate limit: 8 requests per 60 seconds.",
        "Password-reset request does not reveal whether the identifier exists.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("11.1 Configuration", level=2)
    add_table(
        doc,
        ["Setting", "Location", "Notes"],
        [
            ["JWT issuer / audience / key", "appsettings.json → Jwt", "AccessTokenMinutes default 30"],
            ["Database", "ConnectionStrings:DefaultConnection", "PostgreSQL"],
            ["Schema support", "Database:EnsureApiSupportTables", "Roles, approval, school-change, password-reset, rejected_at, last_login_at"],
            ["Public web base URL", "App:PublicWebBaseUrl", "Reset-password email link base"],
            ["Mobile mocks", "USE_MOCKS=true", "Uses MockAuthRepository"],
        ],
        [2800, 3200, 3360],
    )

    # --- 12 ---
    doc.add_heading("12. Client Surfaces (React vs Flutter)", level=1)
    add_note(
        doc,
        "Note",
        "API endpoints are shared. UI coverage differs — see matrix below.",
    )
    add_table(
        doc,
        ["Capability", "React (Web)", "Flutter Mobile"],
        [
            ["login-status / set-initial-password / login", "Yes", "Yes"],
            ["LockedPendingSchoolChange messaging", "Yes — /account-locked", "Yes — in-login locked step"],
            ["Request access / register", "Yes", "Yes"],
            ["Password-reset request", "Yes — /forgot-password", "Yes"],
            ["Password-reset email complete", "Yes — /reset-password", "No (open Web link)"],
            ["Password-reset clear (admin/parent)", "Yes — notification bell", "No (Web admin)"],
            ["Logout with refresh revoke", "Yes", "Yes"],
            ["Multi-role switch", "Yes — Acting as toggle", "Yes — Profile dropdown"],
            ["Request / remove Parent|Teacher", "Yes — /account + role-requests admin", "No / limited (Web first)"],
            ["School/campus change request", "Yes — /account", "Yes — Profile"],
            ["School-change admin apply/reject", "Yes — directory", "No (Web only)"],
            ["Active Directory manage", "Yes", "No (Web only)"],
            ["Full profile edit / deactivate", "Yes — /account", "No / limited"],
        ],
        [3600, 2880, 2880],
    )

    doc.add_heading("12.1 React paths", level=2)
    add_table(
        doc,
        ["Surface", "Path / area", "Behavior"],
        [
            ["Login", "/login", "Multi-step status machine"],
            ["Forgot password", "/forgot-password", "Username (email) → request"],
            ["Reset password", "/reset-password?token=", "Email-token complete"],
            ["Account locked", "/account-locked", "School-change lock messaging"],
            ["Pending registrations", "/admin/directory/registrations", "Scoped activation; reject reason"],
            ["Pending role requests", "/admin/directory/role-requests", "Approve/reject additional Parent|Teacher"],
            ["School changes (admin)", "/admin/directory/school-changes", "Web only"],
            ["Password reset clear", "Notifications bell", "Web only"],
            ["Account / profile", "/account", "Roles Current/Switch/Remove + role request + school change"],
        ],
        [2400, 3000, 3960],
    )

    doc.add_heading("12.2 Flutter Mobile paths", level=2)
    add_table(
        doc,
        ["Surface", "Path / area", "Behavior"],
        [
            ["Login", "/login", "Same status machine; locked is in-page step"],
            ["Profile", "/profile", "Role switch + school-change request"],
            ["Forgot password", "Login sheet", "password-reset request only"],
        ],
        [2400, 3000, 3960],
    )

    # --- 13 ---
    doc.add_heading("13. Related Source Files", level=1)
    add_table(
        doc,
        ["Area", "File"],
        [
            ["API controller", "WebApi/.../Controllers/AuthController.cs"],
            ["Business logic", "WebApi/.../Application/Auth/AuthService.cs"],
            ["Domain user", "WebApi/.../Domain/Auth/User.cs"],
            ["Domain role rules", "WebApi/.../Domain/Auth/UserRoleRules.cs"],
            ["Role request entity", "WebApi/.../Domain/Auth/UserRoleRequest.cs"],
            ["Password reset entity", "WebApi/.../Domain/Auth/UserPasswordResetRequest.cs"],
            ["User repository", "WebApi/.../Repositories/UserRepository.cs"],
            ["Schema initializer", "WebApi/.../ApiSupportSchemaInitializer.cs"],
            ["React auth", "React/src/features/authentication/"],
            ["Flutter auth", "Mobile/lib/features/authentication/"],
            ["Flutter profile", "Mobile/lib/features/profile/.../profile_page.dart"],
            ["QA companion", "docs/03_RankUp_User_Creation_Approval_QA.html"],
            ["This doc", "docs/02_RankUp_Authentication_Logic.html|.docx|.py"],
        ],
        [2400, 6960],
    )

    add_note(
        doc,
        "Version",
        "Aligned with live API (2 Aug 2026). Multi-role Parent+Teacher; role request/remove; "
        "reject reason; Acting as toggle. Client parity: React = full admin/web; "
        "Flutter Mobile = auth + school-change request + role switch; "
        "role-request admin queue and remove-role remain Web-first.",
        "EEF8F1",
    )

    doc.core_properties.title = "RankUp Education — Authentication & Login Logic"
    doc.core_properties.subject = (
        "Authentication, email username, scoped activation, forgot-password first-wins, "
        "app_approval, soft reject, school-change lock, JWT, client parity"
    )
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_doc()
