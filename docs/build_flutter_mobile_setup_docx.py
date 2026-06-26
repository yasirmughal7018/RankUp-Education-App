from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/Flutter_Mobile_Setup_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_table_borders(table, color="B7C6D8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_hyperlinkless_code(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(31, 41, 55)
    return run


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_note_table(doc, label, text, fill, border="B7C6D8"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, border)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(text)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], "E8EEF5")
        p = header_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(22, 57, 87)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if value.startswith("CODE:"):
                add_hyperlinkless_code(p, value.replace("CODE:", "", 1))
            else:
                p.add_run(value)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, "2F3A4A")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "172033")
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(248, 250, 252)
    doc.add_paragraph()


def set_document_styles(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor(23, 50, 77)
    title.paragraph_format.space_after = Pt(3)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(91, 100, 114)
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(46, 116, 181)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.25

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor(46, 116, 181)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.25

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(31, 77, 120)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.25


def build_doc():
    doc = Document()
    set_document_styles(doc)

    doc.add_paragraph("Flutter Mobile App Setup Guide", style="Title")
    doc.add_paragraph("Optimized installation notes for Android and iOS Flutter app development.", style="Subtitle")

    add_table(
        doc,
        ["Prepared for", "Prepared on", "Primary machine", "Purpose"],
        [["RankUp Education workspace", "June 26, 2026", "Windows development PC", "Reuse this checklist before creating another Flutter app"]],
        [2200, 1800, 2200, 3160],
    )

    doc.add_heading("Executive Summary", level=1)
    add_note_table(
        doc,
        "Completed",
        "The optimized Android Flutter setup on this Windows machine is complete and verified. Flutter, Java, Android SDK command-line tools, Android emulator, Android API 36 system image, a Pixel 7 API 36 emulator, VS Code, and the Flutter/Dart VS Code extensions are installed.",
        "EDF8F2",
        "90CBAE",
    )
    add_note_table(
        doc,
        "Important iOS limit",
        "iOS simulator testing and iOS builds cannot be completed on Windows. A Mac with macOS, Xcode, iOS Simulator, CocoaPods, and Flutter is required for the iOS side.",
        "FFF8E8",
        "E5C26F",
    )

    doc.add_heading("Optimized Required Software List", level=1)
    add_table(
        doc,
        ["Software", "Status", "Reason"],
        [
            ["Flutter SDK", "Installed", "Core framework and included Dart SDK for building Flutter apps."],
            ["OpenJDK 17", "Installed", "Required by Android build tooling and Gradle-based Android builds."],
            ["Android SDK command-line tools", "Installed", "Required for sdkmanager, Android platforms, build tools, and emulator packages."],
            ["Android Platform Tools", "Installed", "Provides adb for emulator and device communication."],
            ["Android Build Tools 36.0.0", "Installed", "Required for compiling Android application packages."],
            ["Android Platform 36", "Installed", "Target platform installed for Android API 36 development."],
            ["Android Emulator", "Installed", "Required to run Android virtual devices without installing Android Studio."],
            ["Android API 36 default x86_64 system image", "Installed", "Optimized emulator image selected to reduce disk usage compared with larger Google APIs images."],
            ["VS Code", "Installed", "Lightweight code editor for Flutter development."],
            ["VS Code Flutter and Dart extensions", "Installed", "Provides Flutter commands, Dart analysis, debugging, and editor integration."],
            ["Android Studio IDE", "Not installed intentionally", "Skipped because the optimized setup uses VS Code plus Android SDK command-line tools."],
            ["Separate Dart SDK", "Not installed intentionally", "Skipped because Flutter already includes the required Dart SDK."],
        ],
        [2200, 2100, 5060],
    )

    doc.add_heading("Installed Paths and Versions", level=1)
    add_table(
        doc,
        ["Item", "Path or Version"],
        [
            ["Flutter SDK", "C:\\Users\\yasir\\devtools\\flutter, Flutter 3.44.4 stable, Dart 3.12.2"],
            ["Java", "C:\\Users\\yasir\\devtools\\jdk-17"],
            ["Android SDK", "C:\\Users\\yasir\\devtools\\android-sdk"],
            ["Android emulator", "Android Emulator 36.6.11.0"],
            ["Android virtual device", "Pixel_7_API_36"],
            ["VS Code", "C:\\Users\\yasir\\AppData\\Local\\Programs\\Microsoft VS Code, version 1.126.0"],
            ["VS Code extensions", "Dart-Code.flutter and Dart-Code.dart-code"],
        ],
        [3000, 6360],
    )

    doc.add_heading("Environment Variables", level=1)
    add_table(
        doc,
        ["Variable", "Value"],
        [
            ["JAVA_HOME", "C:\\Users\\yasir\\devtools\\jdk-17"],
            ["ANDROID_HOME", "C:\\Users\\yasir\\devtools\\android-sdk"],
            ["ANDROID_SDK_ROOT", "C:\\Users\\yasir\\devtools\\android-sdk"],
            ["ANDROID_AVD_HOME", "C:\\Users\\yasir\\.android\\avd"],
            ["User Path", "Includes Java, Flutter, Android command-line tools, platform-tools, emulator, and VS Code bin."],
        ],
        [2600, 6760],
    )
    doc.add_paragraph("Open a new terminal, VS Code window, or Cursor window after setup so these variables are loaded.")

    doc.add_heading("Verification Commands", level=1)
    add_code_block(
        doc,
        "flutter doctor -v\nadb version\nemulator -version\nemulator -list-avds\nemulator-check.exe accel",
    )
    doc.add_paragraph("Expected result: Flutter and the Android toolchain should be green. A missing Visual Studio C++ workload warning is not relevant for Android/iOS mobile Flutter work on this Windows machine.")

    doc.add_heading("Running the Android Emulator", level=1)
    doc.add_paragraph("Open a new PowerShell terminal and run:")
    add_code_block(doc, "emulator @Pixel_7_API_36")
    doc.add_paragraph("If the emulator cannot find the AVD before a restart, make sure this environment variable is available in that terminal:")
    add_code_block(doc, '$env:ANDROID_AVD_HOME = "C:\\Users\\yasir\\.android\\avd"')

    doc.add_heading("iOS Development Requirement", level=1)
    doc.add_paragraph("Flutter iOS code can be written on Windows, but the iOS simulator and iOS build chain require macOS. For iOS testing and release builds, prepare a Mac with:")
    for item in [
        "macOS compatible with the current Xcode version.",
        "Xcode installed from the Mac App Store or Apple Developer downloads.",
        "Xcode command-line tools configured with xcode-select.",
        "iOS Simulator runtime installed through Xcode.",
        "CocoaPods installed for Flutter plugins that need native iOS dependencies.",
        "Flutter SDK installed and checked with flutter doctor -v.",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph("Xcode cannot replace the Android SDK for Android builds. It is required for iOS, while Android still needs the Android SDK, platform tools, build tools, platform package, and emulator or physical Android device.")

    doc.add_heading("When to Install a Larger Android System Image", level=1)
    doc.add_paragraph("The current Android emulator uses a smaller default x86_64 image for an optimized setup. Install a Google APIs or Google Play system image only if the app requires emulator-side Google Play Services, Google Maps, Play Billing, Firebase auth flows that depend on Google services, or Play Store testing.")

    doc.add_heading("Reusable Future Setup Prompt", level=1)
    doc.add_paragraph("Use this prompt before setting up another Flutter app environment:")
    add_code_block(
        doc,
        "I need to set up an optimized Flutter development environment for a mobile app that will run on Android and iOS. Please install only required software and avoid heavy optional tools unless they are necessary.\n\n"
        "Before installing anything, first share the exact software list and explain why each item is required or skipped.\n\n"
        "Target machine:\n"
        "- Operating system: [Windows/macOS/Linux]\n"
        "- Editor preference: VS Code\n"
        "- Android testing: emulator and/or physical Android device\n"
        "- iOS testing: iOS Simulator if on macOS, otherwise explain that iOS simulator requires a Mac with Xcode\n\n"
        "Required outcome:\n"
        "- Flutter SDK installed\n"
        "- Java installed only if needed for Android builds\n"
        "- Android SDK command-line tools installed\n"
        "- Android platform tools, build tools, platform package, emulator, and one optimized system image installed\n"
        "- A working Android virtual device created\n"
        "- VS Code installed with Flutter and Dart extensions\n"
        "- Environment variables configured\n"
        "- Licenses accepted\n"
        "- Verification completed with flutter doctor -v\n\n"
        "Important rules:\n"
        "- Do not install Android Studio unless it is specifically needed.\n"
        "- Do not install a separate Dart SDK because Flutter includes Dart.\n"
        "- Do not claim iOS simulator support on Windows.\n"
        "- Use a smaller Android system image unless Google Play Services are required.\n"
        "- After installation, provide installed paths, versions, emulator name, environment variables, and commands to verify or run the emulator."
    )

    doc.add_heading("Future App Creation Checklist", level=1)
    for item in [
        "Open a new terminal so the Flutter and Android paths are available.",
        "Run flutter doctor -v and confirm Android toolchain status.",
        "Start the emulator with emulator @Pixel_7_API_36.",
        "Create the app with flutter create app_name.",
        "Open the project in VS Code.",
        "Run flutter devices and confirm the emulator appears.",
        "Run flutter run for Android testing.",
        "For iOS testing, move the project to a Mac with Xcode and run flutter doctor -v there.",
    ]:
        add_number(doc, item)

    doc.core_properties.title = "Flutter Mobile App Setup Guide"
    doc.core_properties.subject = "Optimized Flutter setup documentation for Android and iOS development"
    doc.core_properties.keywords = "Flutter, Android SDK, VS Code, iOS, Xcode, emulator"
    doc.core_properties.comments = "Generated as a reusable future setup reference."
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
