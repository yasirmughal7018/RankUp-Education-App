"""Generate RankUp Education Authentication & Login Logic DOCX documentation."""

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = "docs/RankUp_Authentication_Logic.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_table_borders(table, color="B7C6D8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
        p = table.rows[0].cells[idx].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(22, 57, 87)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].paragraphs[0].add_run(value)
    doc.add_paragraph()


def add_note(doc, label, text, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    doc.add_paragraph()


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, "2F3A4A")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "172033")
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(248, 250, 252)
    doc.add_paragraph()


def set_document_styles(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor(23, 50, 77)
    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(91, 100, 114)
    for level, size in ((1, 16), (2, 13), (3, 12)):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(46, 116, 181)
        h.font.bold = True


def build_doc():
    doc = Document()
    set_document_styles(doc)

    doc.add_paragraph("RankUp Education — Authentication & Login Logic", style="Title")
    doc.add_paragraph(
        "Web API, database model, JWT tokens, registration workflow, class hierarchy, and request flows.",
        style="Subtitle",
    )

    add_table(
        doc,
        ["Prepared for", "Scope", "Auth model"],
        [["RankUp Education workspace", "WebApi + Flutter Mobile", "JWT access token + hashed refresh token"]],
        [2200, 3200, 3960],
    )

    doc.add_heading("1. Executive Summary", level=1)
    for item in [
        "All authenticated API calls use a Bearer JWT access token.",
        "Login validates username and password against app_users. Only is_active = true with a non-null password_hash may sign in.",
        "Registration creates an inactive app_users row. Admin approval activates the account, sets a password, and creates the role profile.",
        "Five roles: SuperAdmin, SchoolAdmin, Teacher, Student, Parent.",
        "Mobile uses one login screen; role comes from the API response after authentication.",
        "All API responses use ApiResponse<T> with success, message, data, and errors.",
    ]:
        add_bullet(doc, item)

    add_note(
        doc,
        "Important",
        "Extra registration form fields (email, school name, reason) are validated but NOT persisted. "
        "Only username (mobile), display_name, role, requested_at, is_active, and password_hash are stored.",
        "FFF8E8",
    )

    doc.add_heading("2. Solution Architecture", level=1)
    doc.add_heading("2.1 Web API layers", level=2)
    add_table(
        doc,
        ["Layer", "Project", "Responsibility"],
        [
            ["API", "RankUpEducation.Api", "HTTP endpoints, JWT middleware, rate limiting"],
            ["Application", "RankUpEducation.Application", "AuthService, business rules, permissions"],
            ["Contracts", "RankUpEducation.Contracts", "Request/response DTOs"],
            ["Domain", "RankUpEducation.Domain", "User, RefreshToken, UserRole"],
            ["Infrastructure", "RankUpEducation.Infrastructure", "EF Core, JWT, password hashing"],
            ["Common", "RankUpEducation.Common", "CORS, rate limiting, health checks"],
        ],
        [1400, 2800, 5160],
    )

    doc.add_heading("2.2 Mobile layers", level=2)
    add_table(
        doc,
        ["Layer", "Path", "Responsibility"],
        [
            ["Presentation", "features/authentication/presentation/", "LoginPage, AuthController"],
            ["Domain", "features/authentication/domain/", "AuthRepository, entities"],
            ["Data", "features/authentication/data/", "ApiAuthRepository, datasources, models"],
            ["Core", "core/storage/token_store.dart", "Secure token persistence"],
        ],
        [1600, 3600, 4160],
    )

    doc.add_heading("3. Database Model", level=1)
    doc.add_heading("3.1 app_users columns", level=2)
    add_table(
        doc,
        ["Column", "Type", "Purpose"],
        [
            ["id", "BIGINT", "Primary key"],
            ["username", "VARCHAR(50)", "Unique login id (mobile for self-registration)"],
            ["display_name", "VARCHAR(50)", "Full name"],
            ["role", "TEXT", "superadmin | schooladmin | teacher | student | parent"],
            ["password_hash", "TEXT NULL", "PBKDF2 hash; NULL while pending"],
            ["is_active", "BOOLEAN", "Login gate"],
            ["requested_at", "TIMESTAMPTZ", "Registration submission time"],
        ],
        [2200, 1800, 5360],
    )

    add_code(doc, "Pending: is_active = false AND password_hash IS NULL\nActive:   is_active = true  AND password_hash IS NOT NULL")

    doc.add_heading("3.2 Supporting tables", level=2)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["refresh_tokens", "Hashed refresh tokens with expiry and revocation"],
            ["device_sessions", "Device registration and push tokens"],
            ["app_user_students / teachers / parents", "Role profiles with mobile_number on approval"],
        ],
        [3200, 6160],
    )

    doc.add_heading("4. Roles & Permissions", level=1)
    add_table(
        doc,
        ["Role", "DB value", "Sample permissions"],
        [
            ["SuperAdmin", "superadmin", "platform.manage, registration.review"],
            ["SchoolAdmin", "schooladmin", "school.manage, registration.review, quiz.manage"],
            ["Teacher", "teacher", "quiz.create, quiz.assign, attendance.mark"],
            ["Student", "student", "quiz.attempt, dashboard.view"],
            ["Parent", "parent", "child.view, dashboard.view"],
        ],
        [1800, 1800, 5760],
    )

    doc.add_heading("5. API Endpoints", level=1)
    add_table(
        doc,
        ["Method", "Route", "Auth", "Description"],
        [
            ["POST", "/api/auth/login", "Anonymous", "Login with username + password"],
            ["POST", "/api/auth/register", "Anonymous", "Submit registration request"],
            ["GET", "/api/auth/registrations/pending", "Admin", "List pending registrations"],
            ["POST", "/api/auth/registrations/{id}/approve", "Admin", "Activate user + create profile"],
            ["POST", "/api/auth/registrations/{id}/reject", "Admin", "Delete pending request"],
            ["POST", "/api/auth/token/refresh", "Anonymous", "Refresh access token"],
            ["POST", "/api/auth/logout", "Anonymous", "Revoke refresh token"],
            ["GET", "/api/auth/me", "Bearer JWT", "Current user profile"],
            ["POST", "/api/auth/password-reset/request", "Anonymous", "Admin-mediated reset"],
        ],
        [900, 3200, 1200, 4060],
    )

    doc.add_heading("6. Registration Flow", level=1)
    add_code(
        doc,
        "Mobile → POST /api/auth/register\n"
        "  → AuthService.RegisterAccountAsync()\n"
        "  → INSERT app_users (username=mobile, display_name, role,\n"
        "                      is_active=false, password_hash=NULL, requested_at=now())\n\n"
        "Admin → POST /api/auth/registrations/{id}/approve\n"
        "  → INSERT profile row (student | teacher | parent)\n"
        "  → user.Activate(passwordHash)\n"
        "  → is_active = true",
    )

    doc.add_heading("6.1 Approve payload by role", level=2)
    add_table(
        doc,
        ["Role", "Required fields on approve"],
        [
            ["Student", "password, schoolId, campusId, grade, studentRollNumber"],
            ["Teacher", "password, schoolId, campusId, teacherCode"],
            ["Parent", "password (optional cnic)"],
        ],
        [1800, 7560],
    )

    doc.add_heading("7. Login Flow", level=1)
    add_code(
        doc,
        "Mobile → POST /api/auth/login { username, password }\n"
        "  1. UserRepository.GetByUsernameAsync()\n"
        "  2. Attach profile context (schoolId, campusId, profileId)\n"
        "  3. user.EnsureCanLogin() — active + password required\n"
        "  4. PasswordHasher.Verify()\n"
        "  5. IssueRefreshToken() → INSERT refresh_tokens (hashed)\n"
        "  6. JwtTokenService.CreateAccessToken()\n"
        "  → LoginResponse { accessToken, refreshToken, user }",
    )

    doc.add_heading("7.1 Login failure cases", level=2)
    add_table(
        doc,
        ["Condition", "HTTP", "Message"],
        [
            ["Unknown username / wrong password", "401", "Invalid username or password"],
            ["Pending registration", "401", "This account is pending admin approval"],
            ["Inactive account", "401", "This account is not active"],
            ["Validation error", "400", "Errors array in ApiResponse"],
        ],
        [3600, 1200, 4560],
    )

    doc.add_heading("8. Token Refresh & Logout", level=1)
    add_code(
        doc,
        "Refresh: POST /api/auth/token/refresh\n"
        "  → Hash refresh token (SHA-256 hex)\n"
        "  → Lookup refresh_tokens, verify active\n"
        "  → Revoke old token, issue new pair\n\n"
        "Logout: POST /api/auth/logout\n"
        "  → Revoke refresh token row\n"
        "  → Mobile clears TokenStore + local cache",
    )

    doc.add_heading("9. JWT Claims", level=1)
    add_table(
        doc,
        ["Claim", "Source"],
        [
            ["sub / userId", "app_users.id"],
            ["name", "app_users.username"],
            ["role", "UserRole enum"],
            ["permissions", "AuthPermissions.ForRole()"],
            ["profileId", "Profile table id (student_id / teacher_id / parent_id)"],
            ["schoolId / campusId", "From profile when applicable (int)"],
        ],
        [2800, 6560],
    )

    doc.add_heading("10. Class Hierarchy — Web API", level=1)
    add_code(
        doc,
        "Api.Controllers.AuthController\n"
        "Application.Auth.AuthService : IAuthService\n"
        "Application.Auth.AuthMapping / RegistrationMapping\n"
        "Application.Auth.AuthPermissions\n"
        "Infrastructure.Authentication.JwtTokenService : ITokenService\n"
        "Infrastructure.Authentication.PasswordHasher : IPasswordHasher\n"
        "Infrastructure.Authentication.CurrentUserService : ICurrentUserService\n"
        "Infrastructure.Persistence.Repositories.UserRepository : IUserRepository\n"
        "Domain.Auth.User, RefreshToken, DeviceSession, UserRole\n"
        "Contracts.Auth.* (LoginRequest, LoginResponse, RegisterAccountRequest, etc.)",
    )

    doc.add_heading("10.1 User domain methods", level=2)
    add_table(
        doc,
        ["Method", "Purpose"],
        [
            ["CreateRegistrationRequest()", "Factory for pending registration"],
            ["EnsureCanLogin()", "Blocks inactive or passwordless accounts"],
            ["Activate(passwordHash)", "Sets password and is_active = true"],
            ["AttachProfileContext()", "Loads school/campus/profile from profile tables"],
            ["AddRefreshToken()", "Links refresh token to user"],
        ],
        [3600, 5760],
    )

    doc.add_heading("10.2 Class Hierarchy — Mobile", level=1)
    add_code(
        doc,
        "presentation/login_page.dart, auth_controller.dart\n"
        "domain/repositories/auth_repository.dart\n"
        "domain/entities/app_user.dart, auth_session.dart, user_role.dart\n"
        "data/repositories/api_auth_repository.dart\n"
        "data/datasources/auth_remote_datasource.dart\n"
        "data/models/app_user_model.dart, auth_session_model.dart",
    )

    doc.add_heading("11. Security", level=1)
    for item in [
        "Passwords: PBKDF2-SHA256, 100,000 iterations.",
        "Refresh tokens stored as SHA-256 hex only.",
        "Login rate limit: 8 requests per 60 seconds.",
        "Password reset does not reveal whether username exists.",
        "JWT signing key configured in appsettings.json → Jwt:SigningKey.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("12. Configuration", level=1)
    add_table(
        doc,
        ["Setting", "Location", "Notes"],
        [
            ["JWT issuer / audience / key", "appsettings.json → Jwt", "AccessTokenMinutes default 30"],
            ["Database", "ConnectionStrings:DefaultConnection", "PostgreSQL"],
            ["Registration schema", "Database:EnsureApiSupportTables", "Auto-applies on startup"],
            ["Mobile mocks", "USE_MOCKS=true", "Uses MockAuthRepository"],
        ],
        [2800, 3200, 3360],
    )

    doc.add_heading("13. Related Source Files", level=1)
    add_table(
        doc,
        ["Area", "File"],
        [
            ["API controller", "WebApi/src/RankUpEducation.Api/Controllers/AuthController.cs"],
            ["Business logic", "WebApi/src/RankUpEducation.Application/Auth/AuthService.cs"],
            ["JWT service", "WebApi/src/RankUpEducation.Infrastructure/Authentication/JwtTokenService.cs"],
            ["User repository", "WebApi/src/RankUpEducation.Infrastructure/Persistence/Repositories/UserRepository.cs"],
            ["Registration SQL", "docs/RankUpEducation_app_users_registration.sql"],
            ["Mobile login", "Mobile/lib/features/authentication/presentation/pages/login_page.dart"],
            ["Mobile API client", "Mobile/lib/features/authentication/data/datasources/auth_remote_datasource.dart"],
        ],
        [2400, 6960],
    )

    add_note(
        doc,
        "Version",
        "Reflects app_users registration model without JSONB column or account_requests table.",
    )

    doc.core_properties.title = "RankUp Education — Authentication & Login Logic"
    doc.core_properties.subject = "Authentication, login, registration, JWT, and class hierarchy"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
